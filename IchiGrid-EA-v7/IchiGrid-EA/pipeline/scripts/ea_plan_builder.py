#!/usr/bin/env python3
"""
EA Plan Builder v3.0 — Definitive Script
=========================================
Parses ALL .docx plan files → EA_Master_vLLM.xlsx ready for 20 collaborative
AI workers on vLLM. Supports unlimited sections (target: 42,000+).

Supports 5 input formats auto-detected per file:
  1. Tables with 6-7 columns (most common, sections 761+)
  2. Pipe-delimited markdown lines in paragraphs
  3. Paragraph cycling (7 fields per module, no tables)
  4. Narrative section blocks (free-form, sections 1-760)
  5. Ichimoku filters (F01-F44)

Features:
  - SHA-256 persistent memory (plan spec hash + code hash)
  - Auto-fix parasitic section IDs
  - Auto-generate file names for narrative modules
  - Diff engine: detect added/modified/removed modules between runs
  - Path hierarchy, Tags, Status, Priority, AssignedTo for vLLM orchestration
  - Data validation dropdowns (Status, Priority, IA-01→IA-20)
  - 5 sheets: PLAN, HASHMAP, STATS, PROMPT_TEMPLATE, LEGEND
  - Double-click mode: auto-scans script directory

Usage:
  # Double-click: place script next to .docx files, double-click
  # CLI full run:
  python ea_plan_builder.py --docx-dir ./plan_docs --output EA_Master_vLLM.xlsx
  # Dry-run preview:
  python ea_plan_builder.py --docx-dir ./plan_docs --dry-run
  # Single file inspection:
  python ea_plan_builder.py --docx-file ./some_file.docx --inspect
  # Incremental update with diff:
  python ea_plan_builder.py --docx-dir ./plan_docs --output EA_Master_vLLM.xlsx --update

Requirements:
  pip install python-docx openpyxl
"""

import argparse
import hashlib
import os
import re
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx required. Run: pip install python-docx")
    input("Press Enter to exit...")
    sys.exit(1)

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation
except ImportError:
    print("ERROR: openpyxl required. Run: pip install openpyxl")
    input("Press Enter to exit...")
    sys.exit(1)


# ═════════════════════════════════════════════════════════════════════
# DATA MODEL
# ═════════════════════════════════════════════════════════════════════

class Module:
    """Represents one extracted module/section from the plan."""
    __slots__ = [
        'section_id', 'module_name', 'files', 'extensions',
        'function', 'dependencies', 'category', 'source_file',
        'format_type', 'section_hash'
    ]

    def __init__(self, section_id, module_name='', files='', extensions='',
                 function='', dependencies='', category='', source_file='',
                 format_type=''):
        self.section_id = str(section_id).strip()
        self.module_name = module_name.strip()
        self.files = files.strip()
        self.extensions = extensions.strip()
        self.function = function.strip()
        self.dependencies = dependencies.strip()
        self.category = category.strip()
        self.source_file = source_file
        self.format_type = format_type
        self.section_hash = ''

    def compute_hash(self):
        """SHA-256 on normalized Section + Module + Function."""
        raw = f"{self.section_id}|{self.module_name}|{self.function}"
        normalized = re.sub(r'\s+', ' ', raw).strip().lower()
        self.section_hash = hashlib.sha256(normalized.encode('utf-8')).hexdigest()
        return self.section_hash

    def generate_file_names(self):
        """Auto-generate file names from module name if missing."""
        name = self.module_name.strip()
        if not name:
            return

        # Step 1: Embedded .mqh/.mq5 class name in title
        mqh_match = re.search(r'(\w+)\.mqh', name)
        mq5_match = re.search(r'(\w+)\.mq5', name)
        if mqh_match:
            base = mqh_match.group(1)
        elif mq5_match:
            base = mq5_match.group(1)
        else:
            # Step 2: CamelCase IA module name
            camel = re.search(r'(IA[A-Z]\w{5,})', name)
            if camel:
                base = camel.group(1)
            else:
                # Step 3: Clean descriptive text → CamelCase
                clean = re.sub(r'[\U0001F000-\U0001FFFF\u2600-\u27BF\u2700-\u27BF]+', '', name)
                clean = re.sub(r'^\d+[\.\)\s:]+\s*', '', clean)
                clean = re.sub(r'\s*[—–\-:]+\s*.*$', '', clean)
                clean = re.sub(r'\(.*?\)', '', clean)
                clean = re.sub(r'[^a-zA-ZÀ-ÿ0-9\s]', '', clean).strip()
                words = clean.split()[:5]
                if not words:
                    return
                base = ''.join(w.capitalize() for w in words)

        # Remove accents for file name safety
        for old, new in {'é':'e','è':'e','ê':'e','ë':'e','à':'a','â':'a',
                         'ù':'u','û':'u','ô':'o','î':'i','ï':'i','ç':'c',
                         'ü':'u','ö':'o','ä':'a'}.items():
            base = base.replace(old, new).replace(old.upper(), new.upper())
        base = re.sub(r'[^a-zA-Z0-9_]', '', base)
        if len(base) < 3:
            return

        files = [f"{base}.mqh"]
        exts = ['.mqh']
        func_lower = (self.function or '').lower()
        if any(kw in func_lower for kw in ['log','journal','historique','export',
                'rapport','score','stat','index','matrice','carte','suivi']):
            files.append(f"{base}Data.csv"); exts.append('.csv')
        elif any(kw in func_lower for kw in ['config','parametre','preset','profil',
                'etat','sauvegarde','memoire','plan','schema','scenario']):
            files.append(f"{base}Config.json"); exts.append('.json')
        else:
            files.append(f"{base}Data.json"); exts.append('.json')

        self.files = ', '.join(files)
        self.extensions = ', '.join(sorted(set(exts)))

    def is_valid(self):
        return bool(self.section_id and self.module_name)

    def __repr__(self):
        return f"Module({self.section_id}, {self.module_name[:40]})"


# ═════════════════════════════════════════════════════════════════════
# PARSERS (5 formats)
# ═════════════════════════════════════════════════════════════════════

def parse_tables(doc, source_file):
    """Format 1: Standard 6-7 column tables."""
    modules = []
    for table in doc.tables:
        if len(table.columns) < 6:
            continue
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        if 'section' not in header[0]:
            continue
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            sid = cells[0] if cells else ''
            if not sid or not re.match(r'^\d+$', sid):
                continue
            m = Module(
                section_id=sid,
                module_name=cells[1] if len(cells) > 1 else '',
                files=cells[2] if len(cells) > 2 else '',
                extensions=cells[3] if len(cells) > 3 else '',
                function=cells[4] if len(cells) > 4 else '',
                dependencies=cells[5] if len(cells) > 5 else '',
                category=cells[6] if len(cells) > 6 else '',
                source_file=source_file, format_type='table'
            )
            if m.is_valid():
                modules.append(m)
    return modules


def parse_pipe_lines(doc, source_file):
    """Format 2: Pipe-delimited markdown table lines in paragraphs."""
    modules = []
    pattern = re.compile(r'\|\s*(\d{3,5})\s*\|')
    for p in doc.paragraphs:
        txt = p.text.strip()
        if '|' not in txt or not pattern.search(txt):
            continue
        parts = [x.strip() for x in txt.split('|') if x.strip()]
        if len(parts) < 5:
            continue
        sid = None; offset = 0
        for i, part in enumerate(parts):
            if re.match(r'^\d{3,5}$', part):
                sid = part; offset = i; break
        if not sid:
            continue
        m = Module(
            section_id=sid,
            module_name=parts[offset+1] if len(parts) > offset+1 else '',
            files=parts[offset+2] if len(parts) > offset+2 else '',
            extensions=parts[offset+3] if len(parts) > offset+3 else '',
            function=parts[offset+4] if len(parts) > offset+4 else '',
            dependencies=parts[offset+5] if len(parts) > offset+5 else '',
            category=parts[offset+6] if len(parts) > offset+6 else '',
            source_file=source_file, format_type='pipe'
        )
        if m.is_valid():
            modules.append(m)
    return modules


def parse_paragraph_cycle(doc, source_file):
    """Format 3: Paragraphs cycling 7 fields per module (no tables)."""
    modules = []
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    header_indices = []
    for i, txt in enumerate(paras):
        if txt.lower() == 'section' and i + 6 < len(paras):
            next_items = [paras[i + j].lower() for j in range(1, 7)]
            if 'nom du module' in next_items[0] or 'module' in next_items[0]:
                header_indices.append(i)
    if not header_indices:
        return modules
    for hi in header_indices:
        idx = hi + 7
        while idx + 6 < len(paras):
            sid = paras[idx]
            if not re.match(r'^\d{3,5}$', sid):
                break
            m = Module(
                section_id=sid,
                module_name=paras[idx+1] if idx+1 < len(paras) else '',
                files=paras[idx+2] if idx+2 < len(paras) else '',
                extensions=paras[idx+3] if idx+3 < len(paras) else '',
                function=paras[idx+4] if idx+4 < len(paras) else '',
                dependencies=paras[idx+5] if idx+5 < len(paras) else '',
                category=paras[idx+6] if idx+6 < len(paras) else '',
                source_file=source_file, format_type='paragraph_cycle'
            )
            if m.is_valid():
                modules.append(m)
            idx += 7
    return modules


def parse_narrative_sections(doc, source_file):
    """Format 4: Narrative design docs (01_760 style)."""
    modules = []
    current = None; content_lines = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        m = re.match(r'^🧩\s*(?:SECTION\s+)?(\d{1,3})\b[.:\s—–-]*\s*(.*)', txt)
        if m:
            if current and current.is_valid():
                current.generate_file_names()
                modules.append(current)
            sid = m.group(1)
            title = re.sub(r'\s*—\s*IchiGridEA.*$', '', m.group(2).strip())
            current = Module(
                section_id=sid,
                module_name=title[:200] if title else f'Section_{sid}',
                source_file=source_file, format_type='narrative'
            )
            content_lines = []; continue
        if current:
            content_lines.append(txt)
            if len(content_lines) <= 5:
                current.function = ' '.join(content_lines)[:500]
    if current and current.is_valid():
        current.generate_file_names()
        modules.append(current)
    return modules


def parse_ichimoku_filters(doc, source_file):
    """Format 5: Ichimoku 44 filters document."""
    modules = []; current_filter = None; current_data = {}
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for txt in paras:
        m = re.match(r'^(\d{1,2})\.\s+(.+)', txt)
        if m and int(m.group(1)) <= 50:
            if current_filter and current_data:
                fid = f"F{int(current_filter):02d}"
                mod = Module(
                    section_id=fid,
                    module_name=current_data.get('name', ''),
                    function=current_data.get('utility', ''),
                    category=current_data.get('type', 'Ichimoku Filter'),
                    source_file=source_file, format_type='ichimoku_filter'
                )
                mod.files = 'IchimokuSignals.mqh'
                mod.extensions = '.mqh'
                if mod.is_valid():
                    modules.append(mod)
            current_filter = m.group(1)
            current_data = {'name': m.group(2).strip()}
            continue
        if current_filter:
            if txt.startswith('🔄'): current_data['relation'] = txt
            elif txt.startswith('🕓'): current_data['moment'] = txt
            elif txt.startswith('⚙️'): current_data['type'] = txt.replace('⚙️','').replace('Type :','').strip()
            elif txt.startswith('🎯'): current_data['utility'] = txt.replace('🎯','').replace('Utilité :','').strip()
    if current_filter and current_data:
        fid = f"F{int(current_filter):02d}"
        mod = Module(section_id=fid, module_name=current_data.get('name',''),
                     function=current_data.get('utility',''),
                     category=current_data.get('type','Ichimoku Filter'),
                     source_file=source_file, format_type='ichimoku_filter')
        mod.files = 'IchimokuSignals.mqh'; mod.extensions = '.mqh'
        if mod.is_valid(): modules.append(mod)
    return modules


# ═════════════════════════════════════════════════════════════════════
# ORCHESTRATOR
# ═════════════════════════════════════════════════════════════════════

def parse_docx(filepath):
    """Auto-detect format and extract all modules from a single .docx."""
    doc = Document(filepath)
    fname = os.path.basename(filepath)
    all_modules = []

    # Check Ichimoku filters (small dedicated file)
    first_text = ' '.join(p.text for p in doc.paragraphs[:5]).lower()
    is_ichimoku = ('44 filtres ichimoku' in first_text or 'filtres ichimoku professionnels' in first_text)
    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    if is_ichimoku and para_count < 500:
        mods = parse_ichimoku_filters(doc, fname)
        if mods:
            return mods, 'ichimoku_filter'

    # Tables first (most common)
    table_mods = parse_tables(doc, fname)
    if table_mods:
        all_modules.extend(table_mods)
    # Pipe lines (may contain extras not in tables)
    pipe_mods = parse_pipe_lines(doc, fname)
    existing_ids = {m.section_id for m in all_modules}
    for m in pipe_mods:
        if m.section_id not in existing_ids:
            all_modules.append(m); existing_ids.add(m.section_id)
    if all_modules:
        return all_modules, 'table+pipe'

    # Paragraph cycling
    cycle_mods = parse_paragraph_cycle(doc, fname)
    if cycle_mods:
        return cycle_mods, 'paragraph_cycle'

    # Narrative sections
    narr_mods = parse_narrative_sections(doc, fname)
    if narr_mods:
        return narr_mods, 'narrative'

    return all_modules, 'unknown'


def sort_key(sid):
    """Sort key: Fxx first, then numeric."""
    if sid.startswith('F'):
        return (-1, int(sid[1:]))
    try:
        return (0, int(sid))
    except ValueError:
        return (1, 0)


def parse_directory(docx_dir):
    """Parse all .docx files in a directory."""
    all_modules = {}
    stats = {'files': 0, 'formats': {}, 'duplicates': 0}

    docx_files = sorted(Path(docx_dir).glob('*.docx'))
    if not docx_files:
        print(f"  WARNING: No .docx files found in {docx_dir}")
        return [], stats

    for fpath in docx_files:
        if fpath.name.startswith('~'):
            continue
        print(f"  Parsing: {fpath.name}...", end=' ')
        try:
            mods, fmt = parse_docx(str(fpath))
            stats['files'] += 1
            stats['formats'][fmt] = stats['formats'].get(fmt, 0) + 1
            new_count = 0
            for m in mods:
                if m.section_id not in all_modules:
                    m.compute_hash()
                    all_modules[m.section_id] = m
                    new_count += 1
                else:
                    stats['duplicates'] += 1
                    existing = all_modules[m.section_id]
                    if len(m.function) > len(existing.function):
                        m.compute_hash()
                        all_modules[m.section_id] = m
            print(f"{len(mods)} found, {new_count} new [{fmt}]")
        except Exception as e:
            print(f"ERROR: {e}")

    sorted_modules = [all_modules[k] for k in sorted(all_modules.keys(), key=sort_key)]

    # Post-processing: auto-generate file names
    gen = sum(1 for m in sorted_modules if not m.files.strip() and (m.generate_file_names() or True) and m.files.strip())
    if gen:
        print(f"  Auto-generated file names for {gen} modules")

    # Post-processing: fix parasitic IDs
    fixed = 0
    for m in sorted_modules:
        if not m.section_id.isdigit():
            continue
        sid = int(m.section_id)
        fname = m.source_file.replace('.docx', '').replace('-', '_')
        parts = fname.split('_')
        rs = re = None
        for i in range(len(parts) - 1):
            if parts[i].isdigit() and parts[i+1].isdigit():
                a, b = int(parts[i]), int(parts[i+1])
                if b > a and b <= 100000:
                    rs, re = a, b; break
        if rs is None:
            continue
        if sid > re * 2:
            sid_str = m.section_id
            for trim in range(1, len(sid_str) - 3):
                candidate = int(sid_str[trim:])
                if rs <= candidate <= re:
                    m.section_id = str(candidate)
                    m.compute_hash()
                    fixed += 1; break
    if fixed:
        print(f"  Fixed {fixed} parasitic section IDs")
        deduped = {}
        for m in sorted_modules:
            if m.section_id not in deduped:
                deduped[m.section_id] = m
            else:
                if len(m.function) > len(deduped[m.section_id].function):
                    deduped[m.section_id] = m
        sorted_modules = [deduped[k] for k in sorted(deduped.keys(), key=sort_key)]

    return sorted_modules, stats


# ═════════════════════════════════════════════════════════════════════
# vLLM HELPERS (Path, Tags, Status, Priority)
# ═════════════════════════════════════════════════════════════════════

def generate_path(section_id, category):
    """Build hierarchical path from section ID."""
    if section_id.startswith('F'):
        return f"EA::IchimokuSignals::Filter{section_id}"
    if not section_id.isdigit():
        return f"EA::Misc::{section_id}"
    sid = int(section_id)
    if sid <= 100:     return f"EA::Core::S{sid}"
    elif sid <= 760:   return f"EA::Core::Advanced::S{sid}"
    elif sid <= 2860:  return f"EA::L1::S{sid}"
    elif sid <= 5000:  return f"EA::L2::S{sid}"
    elif sid <= 10000: return f"EA::L2::Extended::S{sid}"
    elif sid <= 15000: return f"EA::L3::S{sid}"
    else:              return f"EA::L3::Extended::S{sid}"


def generate_tags(category, function, name):
    """Auto-generate tags from module metadata."""
    tags = set()
    text = f"{category} {function} {name}".lower()
    for keyword, tag in {
        'clone':'clone','shadow':'shadow','recovery':'recovery','fractal':'fractal',
        'entropy':'entropy','fatigue':'fatigue','drawdown':'risk','volatil':'volatility',
        'stress':'stress','trajectory':'trajectory','prediction':'prediction',
        'forecast':'prediction','filter':'filter','signal':'signal','execution':'execution',
        'memory':'memory','cortex':'cortex','dashboard':'ui','grid':'grid',
        'ichimoku':'ichimoku','protection':'protection','export':'io','import':'io',
        'log':'logging','score':'scoring','classif':'classifier','limit':'limiter',
        'mapper':'mapper','clone':'clone','mutation':'mutation','cluster':'cluster',
    }.items():
        if keyword in text:
            tags.add(tag)
    return ', '.join(sorted(tags)[:5]) if tags else ''


def determine_status(fmt, function):
    if fmt == 'narrative': return 'draft'
    return 'spec' if function and len(function) > 30 else 'draft'


def determine_priority(section_id, deps):
    if section_id.startswith('F'): return 'P0'
    if not section_id.isdigit(): return 'P2'
    sid = int(section_id)
    if sid <= 100:   return 'P0'
    elif sid <= 760:  return 'P1'
    elif sid <= 2860: return 'P1'
    else:             return 'P2'


# ═════════════════════════════════════════════════════════════════════
# EXCEL GENERATOR (vLLM Edition — 5 sheets, 18 columns)
# ═════════════════════════════════════════════════════════════════════

# Style palette
C = {
    'dark': '0D1B2A', 'gold': 'FFD700', 'header_bg': '1B2A4A', 'header_fg': 'FFFFFF',
    'border': 'C0C0C0', 'alt': 'FAFAFA', 'green': 'E8F5E9', 'blue': 'E3F2FD',
    'yellow': 'FFF8E1', 'purple': 'F3E5F5', 'orange': 'FFF3E0', 'red': 'FFEBEE',
    'gray': 'F5F5F5', 'stats_hdr': '2E7D32',
}
THIN = Border(left=Side('thin', C['border']), right=Side('thin', C['border']),
              top=Side('thin', C['border']), bottom=Side('thin', C['border']))
H_FONT = Font('Arial', bold=True, size=10, color=C['header_fg'])
H_FILL = PatternFill('solid', fgColor=C['header_bg'])
H_ALIGN = Alignment(horizontal='center', vertical='center', wrap_text=True)
D_FONT = Font('Arial', size=9)
HASH_FONT = Font('Consolas', size=8, color='888888')
WRAP = Alignment(vertical='top', wrap_text=True)
TOP = Alignment(vertical='top')

STATUS_FILLS = {s: PatternFill('solid', fgColor=c) for s, c in
    [('draft',C['gray']),('spec',C['blue']),('ready',C['yellow']),
     ('generating',C['orange']),('review',C['purple']),('done',C['green']),('error',C['red'])]}
PRIO_FILLS = {p: PatternFill('solid', fgColor=c) for p, c in
    [('P0',C['red']),('P1',C['orange']),('P2',C['yellow']),('P3',C['gray'])]}
ALT_FILL = PatternFill('solid', fgColor=C['alt'])


def build_excel(modules, output_path, stats):
    """Generate the complete EA_Master_vLLM.xlsx (5 sheets, 18 columns)."""
    wb = openpyxl.Workbook()
    now = datetime.now()
    now_str = now.strftime('%Y-%m-%d')
    now_full = now.strftime('%Y-%m-%d %H:%M')

    # ── SHEET 1: PLAN (18 columns) ──────────────────────────────────
    ws = wb.active; ws.title = 'PLAN'
    ws.sheet_properties.tabColor = C['header_bg']

    # Title
    ws.merge_cells('A1:R1')
    tc = ws['A1']
    tc.value = f"EA MASTER PLAN — {len(modules)} modules — vLLM Industrial Code Generation — {now_str}"
    tc.font = Font('Arial', bold=True, size=13, color=C['gold'])
    tc.fill = PatternFill('solid', fgColor=C['dark'])
    tc.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 32

    # 18 headers
    headers = [
        ('A','Section',10), ('B','Path',28), ('C','ModuleName',45),
        ('D','Files',40), ('E','Extensions',12), ('F','Function',65),
        ('G','Dependencies',35), ('H','Category',25), ('I','Tags',20),
        ('J','PlanHash',20), ('K','SourceFile',22),
        ('L','Status',12), ('M','Priority',8), ('N','AssignedTo',10),
        ('O','CodeHash',20), ('P','GenDate',12), ('Q','GenModel',15), ('R','Notes',30),
    ]
    for col_letter, name, width in headers:
        cell = ws[f'{col_letter}2']
        cell.value = name; cell.font = H_FONT; cell.fill = H_FILL
        cell.alignment = H_ALIGN; cell.border = THIN
        ws.column_dimensions[col_letter].width = width
    ws.row_dimensions[2].height = 28

    # Write data
    for i, m in enumerate(modules):
        row = i + 3
        path = generate_path(m.section_id, m.category)
        tags = generate_tags(m.category, m.function, m.module_name)
        status = determine_status(m.format_type, m.function)
        priority = determine_priority(m.section_id, m.dependencies)
        short_hash = m.section_hash[:16] + '...' if m.section_hash else ''

        values = [
            m.section_id, path, m.module_name,
            m.files, m.extensions, m.function,
            m.dependencies, m.category, tags,
            short_hash, m.source_file,
            status, priority, '', '', '', '', '',
        ]
        for ci, val in enumerate(values):
            cell = ws.cell(row=row, column=ci+1, value=val)
            cell.border = THIN
            cell.font = HASH_FONT if ci == 9 else D_FONT
            cell.alignment = WRAP if ci in (3,5,6) else TOP
            if ci == 11 and val in STATUS_FILLS:
                cell.fill = STATUS_FILLS[val]
            elif ci == 12 and val in PRIO_FILLS:
                cell.fill = PRIO_FILLS[val]
            elif i % 2 == 0:
                cell.fill = ALT_FILL
        ws.row_dimensions[row].height = 22

    # Freeze + filter
    ws.freeze_panes = 'A3'
    last_row = len(modules) + 2
    ws.auto_filter.ref = f'A2:R{last_row}'

    # Data validations (dropdowns)
    dv_status = DataValidation(type='list', formula1='"draft,spec,ready,generating,review,done,error"', allow_blank=True)
    dv_status.error = 'Valeurs: draft, spec, ready, generating, review, done, error'
    ws.add_data_validation(dv_status); dv_status.add(f'L3:L{last_row}')

    dv_prio = DataValidation(type='list', formula1='"P0,P1,P2,P3"', allow_blank=True)
    ws.add_data_validation(dv_prio); dv_prio.add(f'M3:M{last_row}')

    ia_list = ','.join([f'IA-{i:02d}' for i in range(1, 21)])
    dv_ia = DataValidation(type='list', formula1=f'"{ia_list}"', allow_blank=True)
    ws.add_data_validation(dv_ia); dv_ia.add(f'N3:N{last_row}')

    # ── SHEET 2: HASHMAP (full hashes for machines) ─────────────────
    ws_h = wb.create_sheet('HASHMAP')
    ws_h.sheet_properties.tabColor = '4CAF50'
    for ci, (name, w) in enumerate([('Section',10),('PlanHash_SHA256',68),('CodeHash_SHA256',68),('LastVerified',12)]):
        cell = ws_h.cell(row=1, column=ci+1, value=name)
        cell.font = H_FONT; cell.fill = H_FILL; cell.border = THIN
        ws_h.column_dimensions[get_column_letter(ci+1)].width = w
    for i, m in enumerate(modules):
        ws_h.cell(row=i+2, column=1, value=m.section_id).font = D_FONT
        ws_h.cell(row=i+2, column=2, value=m.section_hash).font = HASH_FONT
        ws_h.cell(row=i+2, column=3, value='').font = HASH_FONT
        ws_h.cell(row=i+2, column=4, value=now_str).font = D_FONT

    # ── SHEET 3: STATS ──────────────────────────────────────────────
    ws_s = wb.create_sheet('STATS')
    ws_s.sheet_properties.tabColor = 'FF9800'

    status_counts = Counter(determine_status(m.format_type, m.function) for m in modules)
    prio_counts = Counter(determine_priority(m.section_id, m.dependencies) for m in modules)
    fmt_counts = Counter(m.format_type for m in modules)
    src_counts = Counter(m.source_file for m in modules)
    tag_counts = Counter()
    for m in modules:
        for t in generate_tags(m.category, m.function, m.module_name).split(','):
            t = t.strip()
            if t: tag_counts[t] += 1

    numeric_ids = [int(m.section_id) for m in modules if m.section_id.isdigit()]
    rows_data = [
        ('EA Master — vLLM Dashboard', '', True),
        (f'Generated: {now_full}', '', False),
        ('', '', False),
        ('OVERVIEW', '', True),
        ('Total Modules', len(modules), False),
        ('Unique SHA-256', len(set(m.section_hash for m in modules if m.section_hash)), False),
        ('Source Files', stats.get('files', len(src_counts)), False),
        ('Section Range', f'{min(numeric_ids)} → {max(numeric_ids)}' if numeric_ids else 'N/A', False),
        ('Duplicates Merged', stats.get('duplicates', 0), False),
        ('', '', False),
        ('BY STATUS', '', True),
    ] + [(s, c, False) for s, c in sorted(status_counts.items())] + [
        ('', '', False), ('BY PRIORITY', '', True),
    ] + [(p, c, False) for p, c in sorted(prio_counts.items())] + [
        ('', '', False), ('BY FORMAT', '', True),
    ] + [(f, c, False) for f, c in sorted(fmt_counts.items())] + [
        ('', '', False), ('vLLM PROGRESS', '', True),
        ('Code Generated', 0, False), ('In Review', 0, False),
        ('Done', 0, False), ('Remaining', len(modules), False),
        ('', '', False), ('TOP TAGS', '', True),
    ] + [(t, c, False) for t, c in tag_counts.most_common(15)] + [
        ('', '', False), ('BY SOURCE FILE', '', True),
    ] + [(s, c, False) for s, c in sorted(src_counts.items())]

    for ri, (a, b, is_header) in enumerate(rows_data):
        ca = ws_s.cell(row=ri+1, column=1, value=a)
        cb = ws_s.cell(row=ri+1, column=2, value=b)
        ca.font = cb.font = Font('Arial', size=10)
        if ri == 0:
            ca.font = Font('Arial', size=14, bold=True, color=C['gold'])
            ca.fill = PatternFill('solid', fgColor=C['dark'])
        elif is_header:
            ca.font = Font('Arial', bold=True, color='FFFFFF')
            ca.fill = PatternFill('solid', fgColor=C['stats_hdr'])
    ws_s.column_dimensions['A'].width = 40
    ws_s.column_dimensions['B'].width = 25

    # ── SHEET 4: PROMPT_TEMPLATE ────────────────────────────────────
    ws_p = wb.create_sheet('PROMPT_TEMPLATE')
    ws_p.sheet_properties.tabColor = '9C27B0'
    ws_p['A1'] = PROMPT_TEMPLATE_TEXT
    ws_p['A1'].font = Font('Consolas', size=10)
    ws_p['A1'].alignment = Alignment(wrap_text=True, vertical='top')
    ws_p.column_dimensions['A'].width = 110
    ws_p.row_dimensions[1].height = 1200

    # ── SHEET 5: LEGEND ─────────────────────────────────────────────
    ws_l = wb.create_sheet('LEGEND')
    ws_l.sheet_properties.tabColor = '607D8B'
    for ri, row_data in enumerate(LEGEND_DATA):
        for ci, val in enumerate(row_data):
            cell = ws_l.cell(row=ri+1, column=ci+1, value=val)
            cell.font = Font('Arial', size=10)
            if ri == 0:
                cell.font = Font('Arial', size=14, bold=True, color=C['gold'])
                cell.fill = PatternFill('solid', fgColor=C['dark'])
            elif val in ('IDENTIFICATION','SPECIFICATION','TRAÇABILITÉ','WORKFLOW vLLM','STATUS VALUES'):
                cell.font = Font('Arial', bold=True, color='FFFFFF')
                cell.fill = PatternFill('solid', fgColor=C['stats_hdr'])
    ws_l.merge_cells('A1:C1')
    ws_l.column_dimensions['A'].width = 20
    ws_l.column_dimensions['B'].width = 45
    ws_l.column_dimensions['C'].width = 65

    wb.save(output_path)
    return output_path


# ═════════════════════════════════════════════════════════════════════
# STATIC CONTENT
# ═════════════════════════════════════════════════════════════════════

PROMPT_TEMPLATE_TEXT = """PROMPT TEMPLATE — vLLM Batch Code Generation (20 AI Workers)
=============================================================

For each module row in PLAN, send this prompt to your vLLM instance.
Replace [PLACEHOLDERS] with values from the corresponding row.

═══════════════════════════════════════════════════════════════
SYSTEM PROMPT (same for all workers):
═══════════════════════════════════════════════════════════════

You are an expert MQL5/MQL4 developer specializing in IchiGridEA
algorithmic trading architecture. Generate production-quality code.
Always include @module and @hash tags for traceability.
All code must compile without errors in MetaTrader 5 Build 4540+.
Follow strict module isolation: one class per .mqh file.

═══════════════════════════════════════════════════════════════
USER PROMPT (per module):
═══════════════════════════════════════════════════════════════

Generate the complete code for this EA module:

## Module Specification
- Section: [Section]
- Path: [Path]
- Module: [ModuleName]
- Files: [Files]
- Function: [Function]
- Dependencies: [Dependencies]
- Category: [Category]
- PlanHash: [PlanHash]

## Requirements
1. Generate ALL listed files
2. .mqh: #ifndef guards, class with constructor/destructor, proper types
3. .json: default schema with documented fields
4. .csv: column headers matching the module's data model
5. Header: // @module: [ModuleName] | @hash: [PlanHash] | @gen: [date]
6. #include all Dependencies
7. Error handling + logging in every public method
8. Wrap code in: // gen:begin:[ModuleName] ... // gen:end:[ModuleName]

## Output
Return ONLY code. One code block per file. No explanations.

═══════════════════════════════════════════════════════════════
vLLM BATCH SCRIPT (Python — ready to run):
═══════════════════════════════════════════════════════════════

```python
#!/usr/bin/env python3
\"\"\"
vLLM Batch Code Generator for IchiGridEA
Usage: python vllm_generate.py --xlsx EA_Master_vLLM.xlsx --url http://localhost:8000/v1
\"\"\"
import argparse, csv, hashlib, json, os, sys, time
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

try:
    from openai import OpenAI
except ImportError:
    sys.exit("pip install openai")
try:
    import openpyxl
except ImportError:
    sys.exit("pip install openpyxl")


SYSTEM_PROMPT = \"\"\"You are an expert MQL5 developer for the IchiGridEA trading system.
Generate production-quality code. Include @module and @hash tags.
All code must compile in MetaTrader 5. One class per .mqh file.\"\"\"


def load_modules(xlsx_path, status_filter='ready'):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb['PLAN']
    modules = []
    for row in range(3, ws.max_row + 1):
        sid = ws.cell(row=row, column=1).value
        status = ws.cell(row=row, column=12).value or ''
        if not sid or status.strip().lower() != status_filter:
            continue
        modules.append({
            'row': row,
            'Section': str(sid),
            'Path': ws.cell(row=row, column=2).value or '',
            'ModuleName': ws.cell(row=row, column=3).value or '',
            'Files': ws.cell(row=row, column=4).value or '',
            'Function': ws.cell(row=row, column=6).value or '',
            'Dependencies': ws.cell(row=row, column=7).value or '',
            'Category': ws.cell(row=row, column=8).value or '',
            'PlanHash': ws.cell(row=row, column=10).value or '',
        })
    return modules


def generate_one(client, model, mod, output_dir):
    prompt = f\"\"\"Generate MQL5 code for:
- Section: {mod['Section']}
- Module: {mod['ModuleName']}
- Files: {mod['Files']}
- Function: {mod['Function']}
- Dependencies: {mod['Dependencies']}
- PlanHash: {mod['PlanHash']}

Generate ALL listed files. Include #ifndef guards, error handling, logging.
Header: // @module: {mod['ModuleName']} | @hash: {mod['PlanHash']}
Return ONLY code blocks, no explanations.\"\"\"

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ],
            max_tokens=8192,
            temperature=0.2,
        )
        code = resp.choices[0].message.content
        code_hash = hashlib.sha256(code.encode()).hexdigest()

        # Save code
        safe_name = mod['ModuleName'][:60].replace(' ', '_')
        fpath = os.path.join(output_dir, f"{mod['Section']}_{safe_name}.mqh")
        with open(fpath, 'w', encoding='utf-8') as f:
            f.write(code)

        return {'status': 'done', 'code_hash': code_hash, 'model': model}
    except Exception as e:
        return {'status': 'error', 'code_hash': '', 'model': model, 'error': str(e)}


def worker(worker_id, modules, vllm_url, model, output_dir):
    client = OpenAI(base_url=vllm_url, api_key="none")
    my_mods = [m for m in modules if m.get('AssignedTo') == f'IA-{worker_id:02d}']
    results = []
    for mod in my_mods:
        print(f"  [IA-{worker_id:02d}] Generating [{mod['Section']}] {mod['ModuleName'][:40]}...")
        result = generate_one(client, model, mod, output_dir)
        results.append((mod, result))
        time.sleep(0.1)
    return results


def main():
    parser = argparse.ArgumentParser(description='vLLM Batch Generator')
    parser.add_argument('--xlsx', required=True, help='EA_Master_vLLM.xlsx path')
    parser.add_argument('--url', default='http://localhost:8000/v1', help='vLLM API URL')
    parser.add_argument('--model', default='codellama/CodeLlama-34b-Instruct-hf')
    parser.add_argument('--workers', type=int, default=20, help='Number of AI workers')
    parser.add_argument('--output-dir', default='./generated', help='Output directory')
    args = parser.parse_args()

    os.makedirs(args.output_dir, exist_ok=True)
    modules = load_modules(args.xlsx)
    if not modules:
        print("No modules with status 'ready'. Set Status=ready in Excel first.")
        return

    # Assign round-robin
    for i, mod in enumerate(modules):
        mod['AssignedTo'] = f'IA-{(i % args.workers) + 1:02d}'

    print(f"Generating {len(modules)} modules with {args.workers} workers...")

    all_results = []
    with ThreadPoolExecutor(max_workers=args.workers) as executor:
        futures = [executor.submit(worker, w+1, modules, args.url, args.model, args.output_dir)
                   for w in range(args.workers)]
        for future in as_completed(futures):
            all_results.extend(future.result())

    # Update Excel
    wb = openpyxl.load_workbook(args.xlsx)
    ws = wb['PLAN']
    ws_hash = wb['HASHMAP']
    for mod, result in all_results:
        row = mod['row']
        ws.cell(row=row, column=12).value = result['status']
        ws.cell(row=row, column=14).value = mod['AssignedTo']
        ws.cell(row=row, column=15).value = result['code_hash'][:16] + '...' if result['code_hash'] else ''
        ws.cell(row=row, column=16).value = datetime.now().strftime('%Y-%m-%d')
        ws.cell(row=row, column=17).value = result['model']
        if result.get('error'):
            ws.cell(row=row, column=18).value = result['error'][:200]
        # Update HASHMAP
        for hr in range(2, ws_hash.max_row + 1):
            if str(ws_hash.cell(row=hr, column=1).value) == mod['Section']:
                ws_hash.cell(row=hr, column=3).value = result['code_hash']
                ws_hash.cell(row=hr, column=4).value = datetime.now().strftime('%Y-%m-%d')
                break

    wb.save(args.xlsx)
    done = sum(1 for _, r in all_results if r['status'] == 'done')
    errors = sum(1 for _, r in all_results if r['status'] == 'error')
    print(f"\\nComplete: {done} done, {errors} errors. Excel updated.")


if __name__ == '__main__':
    main()
```

═══════════════════════════════════════════════════════════════
WORKFLOW RÉSUMÉ:
═══════════════════════════════════════════════════════════════

1. Lancer ea_plan_builder.py → génère EA_Master_vLLM.xlsx
2. Dans Excel: passer les modules voulus de "spec" → "ready"
3. Lancer vllm_generate.py → 20 IA génèrent le code en parallèle
4. Status passe à "done" ou "error" automatiquement
5. CodeHash + GenDate + GenModel remplis automatiquement
6. Relancer ea_plan_builder.py --update → diff report des changements
7. Ajouter de nouveaux .docx → relancer → les nouveaux modules apparaissent
"""

LEGEND_DATA = [
    ['EA_Master_vLLM.xlsx — Légende complète', '', ''],
    ['', '', ''],
    ['IDENTIFICATION', '', ''],
    ['Section', 'ID unique du module', 'Numérique (1-42000+) ou Fxx pour filtres Ichimoku'],
    ['Path', 'Chemin hiérarchique', 'EA::Level::Section. Navigation arborescente.'],
    ['ModuleName', 'Nom de la classe MQL5', 'CamelCase. Préfixe IA pour modules IA.'],
    ['', '', ''],
    ['SPECIFICATION', '', ''],
    ['Files', 'Fichiers à générer', '.mqh (code), .json (config), .csv (data)'],
    ['Extensions', 'Types de fichiers', 'Dérivé automatiquement de Files'],
    ['Function', 'Rôle / comportement IA', 'Spec pour la génération de code.'],
    ['Dependencies', 'Modules requis', '#include dans le code généré.'],
    ['Category', 'Bloc thématique', 'Classification fonctionnelle.'],
    ['Tags', 'Tags auto-générés', 'Filtrage: clone, risk, signal, memory...'],
    ['', '', ''],
    ['TRAÇABILITÉ', '', ''],
    ['PlanHash', 'SHA-256 de la spec', 'normalize(Section|ModuleName|Function). Change si spec modifiée.'],
    ['SourceFile', 'Fichier .docx source', 'Traçabilité document original.'],
    ['', '', ''],
    ['WORKFLOW vLLM', '', ''],
    ['Status', 'État du module', 'draft → spec → ready → generating → review → done | error'],
    ['Priority', 'Priorité de build', 'P0=bloquant, P1=core, P2=standard, P3=optionnel'],
    ['AssignedTo', 'IA assignée', 'IA-01 à IA-20. Round-robin ou par spécialisation.'],
    ['CodeHash', 'SHA-256 du code généré', 'Rempli après génération. Détecte si code à jour.'],
    ['GenDate', 'Date de génération', 'Quand le code a été produit.'],
    ['GenModel', 'Modèle vLLM utilisé', 'Ex: codellama-34b, deepseek-coder-v2...'],
    ['Notes', 'Remarques', 'Erreurs, feedback, TODO...'],
    ['', '', ''],
    ['STATUS VALUES', '', ''],
    ['draft', 'Spec incomplète', 'Function trop courte ou vague.'],
    ['spec', 'Spec complète', 'Function décrit le comportement. Prêt pour review.'],
    ['ready', 'Prêt pour génération', 'Validé. Peut être envoyé au vLLM.'],
    ['generating', 'En cours', 'Le vLLM travaille dessus.'],
    ['review', 'En revue', 'Code produit, en validation.'],
    ['done', 'Terminé', 'Code final, hash figé.'],
    ['error', 'Erreur', 'Le vLLM a échoué. Voir Notes.'],
]


# ═════════════════════════════════════════════════════════════════════
# PERSISTENT MEMORY — DIFF ENGINE
# ═════════════════════════════════════════════════════════════════════

def load_existing_hashes(filepath):
    """Load existing Excel and return {section_id: hash}."""
    wb = openpyxl.load_workbook(filepath, data_only=True)
    if 'PLAN' not in wb.sheetnames:
        return {}
    ws = wb['PLAN']
    existing = {}
    for row in range(3, ws.max_row + 1):
        sid = ws.cell(row=row, column=1).value
        sha = ws.cell(row=row, column=10).value  # PlanHash (short)
        name = ws.cell(row=row, column=3).value   # ModuleName
        if sid and sha:
            existing[str(sid)] = {'hash': sha, 'name': name or ''}
    # Also try HASHMAP for full hashes
    if 'HASHMAP' in wb.sheetnames:
        ws_h = wb['HASHMAP']
        for row in range(2, ws_h.max_row + 1):
            sid = ws_h.cell(row=row, column=1).value
            full_hash = ws_h.cell(row=row, column=2).value
            if sid and full_hash and str(sid) in existing:
                existing[str(sid)]['full_hash'] = full_hash
    return existing


def diff_with_existing(excel_path, new_modules):
    """Compare new parse with existing Excel."""
    existing = load_existing_hashes(excel_path)
    new_map = {m.section_id: m for m in new_modules}

    diff = {
        'added': [], 'removed': [], 'modified': [], 'unchanged': [],
        'timestamp': datetime.now().isoformat(),
        'previous_count': len(existing), 'new_count': len(new_modules),
    }

    for sid, mod in new_map.items():
        if sid not in existing:
            diff['added'].append({'section_id': sid, 'name': mod.module_name, 'hash': mod.section_hash})
        else:
            old_hash = existing[sid].get('full_hash', existing[sid]['hash'])
            # Compare full hash if available, otherwise short hash
            if old_hash == mod.section_hash or (len(old_hash) < 64 and mod.section_hash.startswith(old_hash.replace('...',''))):
                diff['unchanged'].append(sid)
            else:
                diff['modified'].append({
                    'section_id': sid, 'name': mod.module_name,
                    'old_hash': old_hash, 'new_hash': mod.section_hash,
                })

    for sid, data in existing.items():
        if sid not in new_map:
            diff['removed'].append({'section_id': sid, 'name': data['name'], 'hash': data.get('full_hash','')})

    return diff


def print_diff_report(diff):
    """Print diff to console."""
    print(f"\n{'='*70}")
    print(f"PERSISTENT MEMORY — DIFF REPORT")
    print(f"{'='*70}")
    print(f"  Previous: {diff['previous_count']} modules")
    print(f"  Current:  {diff['new_count']} modules")
    print(f"  ✅ Unchanged: {len(diff['unchanged'])}")
    print(f"  ➕ Added:     {len(diff['added'])}")
    print(f"  ✏️  Modified:  {len(diff['modified'])}")
    print(f"  ❌ Removed:   {len(diff['removed'])}")
    for label, key, limit in [('➕ NEW', 'added', 15), ('✏️  MODIFIED', 'modified', 10), ('❌ REMOVED', 'removed', 10)]:
        items = diff[key]
        if items:
            print(f"\n  {label} ({len(items)}):")
            for item in items[:limit]:
                print(f"    [{item['section_id']:>5s}] {item['name'][:55]}")
            if len(items) > limit:
                print(f"    ... +{len(items)-limit} more")


def write_diff_report(diff, output_path):
    """Write diff to markdown."""
    lines = [
        f"# EA Master — Diff Report", '',
        f"**Generated**: {diff['timestamp']}",
        f"**Previous**: {diff['previous_count']} | **Current**: {diff['new_count']}", '',
        f"| Status | Count |", f"|--------|-------|",
        f"| ✅ Unchanged | {len(diff['unchanged'])} |",
        f"| ➕ Added | {len(diff['added'])} |",
        f"| ✏️ Modified | {len(diff['modified'])} |",
        f"| ❌ Removed | {len(diff['removed'])} |", '',
    ]
    for label, key in [('➕ Added', 'added'), ('✏️ Modified', 'modified'), ('❌ Removed', 'removed')]:
        if diff[key]:
            lines += [f"## {label} ({len(diff[key])})", '',
                      f"| Section | Module |", f"|---------|--------|"]
            for item in diff[key]:
                lines.append(f"| {item['section_id']} | {item['name'][:60]} |")
            lines.append('')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f"  📄 Diff report: {output_path}")


# ═════════════════════════════════════════════════════════════════════
# CLI
# ═════════════════════════════════════════════════════════════════════

def inspect_file(filepath):
    """Inspect a single file."""
    print(f"\n{'='*70}")
    print(f"INSPECTING: {filepath}")
    doc = Document(filepath)
    print(f"  Paragraphs: {len(doc.paragraphs)} | Tables: {len(doc.tables)}")
    mods, fmt = parse_docx(filepath)
    print(f"  Format: {fmt} | Modules: {len(mods)}")
    if mods:
        for m in mods[:5]:
            print(f"    [{m.section_id}] {m.module_name[:60]}")
        if len(mods) > 5:
            print(f"    ... +{len(mods)-5} more")
        numeric = [int(m.section_id) for m in mods if m.section_id.isdigit()]
        if numeric:
            print(f"  Range: {min(numeric)} → {max(numeric)}")


def main():
    parser = argparse.ArgumentParser(
        description='EA Plan Builder v3.0 — .docx → EA_Master_vLLM.xlsx (20 AI workers)')
    parser.add_argument('--docx-dir', help='Directory containing .docx plan files')
    parser.add_argument('--docx-file', help='Single .docx file to inspect')
    parser.add_argument('--output', default='EA_Master_vLLM.xlsx', help='Output Excel (default: EA_Master_vLLM.xlsx)')
    parser.add_argument('--dry-run', action='store_true', help='Preview without writing')
    parser.add_argument('--inspect', action='store_true', help='Inspect single file')
    parser.add_argument('--update', action='store_true', help='Incremental update + diff')
    args = parser.parse_args()

    if args.docx_file and args.inspect:
        inspect_file(args.docx_file)
        return

    if args.docx_file and not args.docx_dir:
        args.docx_dir = os.path.dirname(args.docx_file) or '.'

    if not args.docx_dir:
        # AUTO MODE: scan script directory
        script_dir = os.path.dirname(os.path.abspath(__file__))
        docx_count = len(list(Path(script_dir).glob('*.docx')))
        if docx_count > 0:
            print(f"[AUTO] Found {docx_count} .docx files in script directory")
            print(f"[AUTO] Using: {script_dir}")
            args.docx_dir = script_dir
            if args.output == 'EA_Master_vLLM.xlsx':
                args.output = os.path.join(script_dir, 'EA_Master_vLLM.xlsx')
        else:
            parser.print_help()
            print(f"\nERROR: No .docx files found in {script_dir}")
            input("\nPress Enter to exit...")
            sys.exit(1)

    print(f"\n{'='*70}")
    print(f"EA PLAN BUILDER v3.0 — vLLM Edition")
    print(f"Scanning: {args.docx_dir}")
    print(f"{'='*70}\n")

    modules, stats = parse_directory(args.docx_dir)

    print(f"\n{'─'*70}")
    print(f"RESULTS:")
    print(f"  Files parsed:     {stats['files']}")
    print(f"  Total modules:    {len(modules)}")
    print(f"  Duplicates merged:{stats['duplicates']}")
    print(f"  Formats:          {stats['formats']}")

    if args.dry_run:
        print(f"\n[DRY-RUN] Would write {len(modules)} modules to {args.output}")
        for m in modules[:10]:
            print(f"  [{m.section_id:>5s}] {m.module_name[:50]:50s} | {m.format_type}")
        if len(modules) > 10:
            print(f"  ... ({len(modules)-10} more)")
        if len(sys.argv) == 1:
            input("\nPress Enter to exit...")
        return

    if not modules:
        print("WARNING: No modules extracted. Check .docx files.")
        if len(sys.argv) == 1:
            input("\nPress Enter to exit...")
        return

    # Diff with existing
    if args.update and os.path.exists(args.output):
        diff = diff_with_existing(args.output, modules)
        print_diff_report(diff)
        write_diff_report(diff, args.output.replace('.xlsx', '_diff.md'))

    print(f"\nGenerating {args.output}...")
    build_excel(modules, args.output, stats)
    print(f"\n✅ Done! {len(modules)} modules → {args.output}")
    print(f"   Sheets: PLAN (18 cols) | HASHMAP | STATS | PROMPT_TEMPLATE | LEGEND")
    print(f"   Ready for vLLM with 20 AI workers.")

    # Keep console open on double-click
    if len(sys.argv) == 1:
        input("\nPress Enter to exit...")


if __name__ == '__main__':
    main()
